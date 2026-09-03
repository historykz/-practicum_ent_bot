"""
Раздел «Мои ошибки» + скачивание файла ошибок + результаты по тестам.

Ошибки собираются из двух источников:
- attempt_answers (обычные тесты уроков сайта);
- live_answers (Live-викторины).

Ошибка группируется по вопросу. Статус:
- «Исправлена» — есть более поздний правильный ответ на этот же вопрос;
- «На повторении» — отвечал повторно, но снова неверно;
- «Не изучено» — после ошибки к вопросу не возвращался.
"""
import asyncio
import csv
import random
import io
import json
from datetime import datetime

import aiohttp_jinja2
from aiohttp import web

import database as db
import utils
from webapp import auth


async def _login(request):
    tg_id = await auth.get_logged_in_tg_id(request)
    if tg_id is None:
        raise web.HTTPFound("/?error=login_required")
    return tg_id


def _answers_history_sync(tg_id: int):
    """Все ответы пользователя (верные и неверные) из обоих источников.
    Возвращает список dict: question_id, correct(0/1), ts, source, test_title,
    subject, topic, qtext, user_ans, correct_ans."""
    user = utils.get_user_by_tg(tg_id)
    user_id = user["id"] if user else None
    rows = []

    # --- Обычные тесты уроков ---
    if user_id:
        for r in db.fetchall(
            """SELECT aa.question_id AS qid, aa.is_correct AS ok, aa.created_at AS ts,
                      q.text AS qtext, q.topic AS topic,
                      uo.text AS user_ans,
                      (SELECT text FROM question_options WHERE question_id=q.id AND is_correct=1 LIMIT 1) AS correct_ans,
                      t.title AS test_title,
                      COALESCE(s.title, '') AS subject
               FROM attempt_answers aa
               JOIN test_attempts ta ON ta.id = aa.attempt_id
               JOIN questions q ON q.id = aa.question_id
               JOIN tests t ON t.id = ta.test_id
               LEFT JOIN question_options uo ON uo.id = aa.selected_option_id
               LEFT JOIN lessons l ON l.test_id = t.id
               LEFT JOIN sections sec ON sec.id = l.section_id
               LEFT JOIN subjects s ON s.id = sec.subject_id
               WHERE ta.user_id=? AND COALESCE(aa.skipped,0)=0""",
            (user_id,),
        ):
            r = dict(r)
            r["source"] = "Тест урока"
            rows.append(r)

    # --- Live-викторины ---
    for r in db.fetchall(
        """SELECT la.question_id AS qid, la.is_correct AS ok, la.created_at AS ts,
                  q.text AS qtext, q.topic AS topic,
                  uo.text AS user_ans,
                  (SELECT text FROM question_options WHERE question_id=q.id AND is_correct=1 LIMIT 1) AS correct_ans,
                  t.title AS test_title, '' AS subject
           FROM live_answers la
           JOIN live_rooms room ON room.id = la.room_id
           JOIN questions q ON q.id = la.question_id
           JOIN tests t ON t.id = room.test_id
           LEFT JOIN question_options uo ON uo.id = la.option_id
           WHERE la.tg_id=?""",
        (tg_id,),
    ):
        r = dict(r)
        r["source"] = "Live-тест"
        rows.append(r)
    return rows


def build_mistakes_sync(tg_id: int) -> dict:
    """Сводка ошибок пользователя: список уникальных ошибок со статусами + статистика."""
    history = _answers_history_sync(tg_id)
    total_answers = len(history)
    total_correct = sum(1 for h in history if h["ok"])
    total_wrong = total_answers - total_correct

    # группируем по вопросу
    by_q = {}
    for h in history:
        q = by_q.setdefault(h["qid"], {"wrong": [], "correct": [], "meta": h})
        (q["correct"] if h["ok"] else q["wrong"]).append(h["ts"] or "")

    mistakes = []
    topic_stats = {}
    for qid, d in by_q.items():
        if not d["wrong"]:
            continue  # по этому вопросу ошибок не было
        m = d["meta"]
        last_wrong = max(d["wrong"])
        correct_after = [c for c in d["correct"] if c and c > last_wrong]
        if correct_after:
            status = "Исправлена"
        elif d["correct"] or len(d["wrong"]) > 1:
            status = "На повторении"
        else:
            status = "Не изучено"
        topic = m.get("topic") or "Без темы"
        ts = topic_stats.setdefault(topic, {"wrong": 0, "total": 0})
        ts["wrong"] += len(d["wrong"])
        mistakes.append({
            "qid": qid, "qtext": m.get("qtext") or "",
            "user_ans": m.get("user_ans") or "—",
            "correct_ans": m.get("correct_ans") or "—",
            "test_title": m.get("test_title") or "",
            "subject": m.get("subject") or "",
            "topic": topic, "source": m.get("source"),
            "wrong_count": len(d["wrong"]),
            "last_wrong": (last_wrong or "")[:16].replace("T", " "),
            "status": status,
        })

    # тема: всего ответов по теме (для %)
    for h in history:
        topic = h.get("topic") or "Без темы"
        topic_stats.setdefault(topic, {"wrong": 0, "total": 0})["total"] += 1

    weak_topics = []
    for topic, s in topic_stats.items():
        if s["total"] == 0:
            continue
        pct = round(s["wrong"] / s["total"] * 100)
        weak_topics.append({"topic": topic, "wrong": s["wrong"], "total": s["total"], "percent": pct})
    weak_topics.sort(key=lambda x: -x["percent"])

    # сортируем: сначала частые и неисправленные
    order = {"Не изучено": 0, "На повторении": 1, "Исправлена": 2}
    mistakes.sort(key=lambda m: (order.get(m["status"], 3), -m["wrong_count"]))

    return {
        "mistakes": mistakes,
        "total_answers": total_answers,
        "total_correct": total_correct,
        "total_wrong": total_wrong,
        "error_percent": round(total_wrong / total_answers * 100) if total_answers else 0,
        "correct_percent": round(total_correct / total_answers * 100) if total_answers else 0,
        "unique_errors": len(mistakes),
        "unfixed": sum(1 for m in mistakes if m["status"] != "Исправлена"),
        "weak_topics": weak_topics[:12],
    }


async def mistakes_page(request):
    tg_id = await _login(request)
    data = await asyncio.to_thread(build_mistakes_sync, tg_id)
    data.update(await auth.nav_context(request))
    return aiohttp_jinja2.render_template("mistakes.html", request, data)


async def mistakes_csv(request):
    # Скачивание файла ошибок — только для админов (ученик смотрит на сайте)
    from webapp import learning
    await learning._require_admin(request)
    tg_id = await _login(request)
    data = await asyncio.to_thread(build_mistakes_sync, tg_id)
    buf = io.StringIO()
    buf.write("﻿")  # BOM — Excel корректно откроет кириллицу
    w = csv.writer(buf, delimiter=";")
    w.writerow(["Вопрос", "Ваш ответ", "Правильный ответ", "Тема", "Тест",
                "Источник", "Ошибок раз", "Последняя ошибка", "Статус"])
    for m in data["mistakes"]:
        w.writerow([m["qtext"], m["user_ans"], m["correct_ans"], m["topic"],
                    m["test_title"], m["source"], m["wrong_count"], m["last_wrong"], m["status"]])
    return web.Response(
        body=buf.getvalue().encode("utf-8"),
        headers={"Content-Type": "text/csv; charset=utf-8",
                 "Content-Disposition": 'attachment; filename="moi_oshibki.csv"'})


async def mistakes_docx(request):
    from webapp import learning
    await learning._require_admin(request)
    tg_id = await _login(request)
    data = await asyncio.to_thread(build_mistakes_sync, tg_id)
    user = await asyncio.to_thread(utils.get_user_by_tg, tg_id)

    def _build():
        import docx
        doc = docx.Document()
        who = (user.get("username") and "@" + user["username"]) or (user.get("first_name") if user else "") or f"id{tg_id}"
        doc.add_heading("Мои ошибки — Smart ENT", 0)
        doc.add_paragraph(f"Пользователь: {who}")
        doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        doc.add_paragraph(
            f"Всего ответов: {data['total_answers']} · верно: {data['total_correct']} "
            f"({data['correct_percent']}%) · ошибок: {data['total_wrong']} ({data['error_percent']}%)")
        if data["weak_topics"]:
            doc.add_heading("Темы для повторения", level=1)
            for t in data["weak_topics"]:
                doc.add_paragraph(f"• {t['topic']} — ошибок {t['wrong']} из {t['total']} ({t['percent']}%)")
        doc.add_heading("Список ошибок", level=1)
        for i, m in enumerate(data["mistakes"], 1):
            doc.add_heading(f"{i}. {m['qtext']}", level=2)
            doc.add_paragraph(f"Ваш ответ: {m['user_ans']}")
            doc.add_paragraph(f"Правильный ответ: {m['correct_ans']}")
            doc.add_paragraph(f"Тема: {m['topic']} · Тест: {m['test_title']} · {m['source']}")
            doc.add_paragraph(f"Ошибок: {m['wrong_count']} · Статус: {m['status']}")
        out = io.BytesIO()
        doc.save(out)
        return out.getvalue()

    try:
        body = await asyncio.to_thread(_build)
    except Exception as e:
        return web.Response(text=f"Не удалось сформировать Word: {e}", status=500)
    return web.Response(
        body=body,
        headers={"Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                 "Content-Disposition": 'attachment; filename="moi_oshibki.docx"'})


# ============ Результаты по тесту (админ): попытки, баллы, время ============

def test_results_sync(test_id: int) -> dict:
    test = db.fetchone("SELECT * FROM tests WHERE id=?", (test_id,))
    if not test:
        return None
    rows = db.fetchall(
        """SELECT ta.id, ta.correct_answers, ta.wrong_answers, ta.skipped_answers,
                  ta.status, ta.start_time, ta.end_time, ta.created_at, ta.attempt_num,
                  u.username, u.first_name, u.tg_id
           FROM test_attempts ta JOIN users u ON u.id = ta.user_id
           WHERE ta.test_id=? AND ta.status IN ('finished','aborted')
             AND COALESCE(ta.publication_id, 0)=0
           ORDER BY ta.id DESC LIMIT 300""",
        (test_id,))
    attempts = []
    for r in rows:
        total = (r["correct_answers"] or 0) + (r["wrong_answers"] or 0) + (r["skipped_answers"] or 0)
        pct = round((r["correct_answers"] or 0) / total * 100) if total else 0
        dur = ""
        if r["start_time"] and r["end_time"]:
            try:
                d = (datetime.fromisoformat(r["end_time"]) - datetime.fromisoformat(r["start_time"])).total_seconds()
                dur = f"{int(d // 60)}м {int(d % 60)}с"
            except (ValueError, TypeError):
                pass
        who = (r["username"] and "@" + r["username"]) or r["first_name"] or f"id{r['tg_id']}"
        attempts.append({
            "who": who, "tg_id": r["tg_id"],
            "correct": r["correct_answers"] or 0, "wrong": r["wrong_answers"] or 0,
            "skipped": r["skipped_answers"] or 0, "percent": pct,
            "attempt_num": r["attempt_num"], "status": r["status"],
            "duration": dur,
            "when": (r["created_at"] or "")[:16].replace("T", " "),
        })
    return {"test": dict(test), "attempts": attempts, "count": len(attempts)}


async def admin_test_results(request):
    from webapp import learning
    await learning._require_admin(request)
    test_id = int(request.match_info["test_id"])
    data = await asyncio.to_thread(test_results_sync, test_id)
    if data is None:
        raise web.HTTPNotFound()
    data.update(await auth.nav_context(request))
    return aiohttp_jinja2.render_template("test_results.html", request, data)


async def admin_test_results_csv(request):
    from webapp import learning
    await learning._require_admin(request)
    test_id = int(request.match_info["test_id"])
    data = await asyncio.to_thread(test_results_sync, test_id)
    if data is None:
        raise web.HTTPNotFound()
    buf = io.StringIO()
    buf.write("﻿")
    w = csv.writer(buf, delimiter=";")
    w.writerow(["Ученик", "ID", "Верно", "Неверно", "Пропущено", "%", "Попытка",
                "Длительность", "Статус", "Когда"])
    for a in data["attempts"]:
        w.writerow([a["who"], a["tg_id"], a["correct"], a["wrong"], a["skipped"],
                    a["percent"], a["attempt_num"], a["duration"], a["status"], a["when"]])
    return web.Response(
        body=buf.getvalue().encode("utf-8"),
        headers={"Content-Type": "text/csv; charset=utf-8",
                 "Content-Disposition": f'attachment; filename="results_test{test_id}.csv"'})


# ===================== Работа над ошибками: карточки и тест =====================

def practice_questions_sync(tg_id: int, limit: int = 60) -> list:
    """Вопросы, в которых ученик ошибался и ещё не исправил — с вариантами.
    Скачать их нельзя, но прогнать карточками или тестом можно сколько угодно."""
    data = build_mistakes_sync(tg_id)
    out = []
    for m in data.get("mistakes", []):
        if m.get("fixed"):
            continue
        opts = db.fetchall(
            "SELECT id, text, is_correct FROM question_options WHERE question_id=? "
            "ORDER BY id", (m["qid"],))
        opts = [dict(o) for o in opts]
        if not opts:
            continue
        correct = next((o for o in opts if o["is_correct"]), None)
        if not correct:
            continue
        random.shuffle(opts)
        out.append({
            "question_id": m["qid"],
            "qtext": m.get("qtext") or "",
            "topic": m.get("topic") or "",
            "test_title": m.get("test_title") or "",
            "options": [{"id": o["id"], "text": o["text"],
                         "correct": bool(o["is_correct"])} for o in opts],
            "correct_text": correct["text"],
        })
        if len(out) >= limit:
            break
    return out


async def mistakes_cards(request):
    """Карточки по своим ошибкам: вопрос → переворот → правильный ответ."""
    tg_id = await _login(request)
    items = await asyncio.to_thread(practice_questions_sync, tg_id)
    ctx = await auth.nav_context(request)
    ctx.update({"items": items, "mode": "cards"})
    return aiohttp_jinja2.render_template("mistakes_practice.html", request, ctx)


async def mistakes_retest(request):
    """Тест по своим ошибкам: те же вопросы, ответы проверяются на месте."""
    tg_id = await _login(request)
    items = await asyncio.to_thread(practice_questions_sync, tg_id)
    ctx = await auth.nav_context(request)
    ctx.update({"items": items, "mode": "test"})
    return aiohttp_jinja2.render_template("mistakes_practice.html", request, ctx)


# ===================== Админ: самые частые ошибки всех учеников ===============

def top_mistakes_sync(limit: int = 300) -> dict:
    """Вопросы, на которых чаще всего ошибаются — худшие сверху.

    Считаем по всем ученикам сразу: сколько раз отвечали, сколько раз мимо и
    какой процент. Это и есть подсказка, что объяснять на уроке заново.
    """
    rows = db.fetchall(
        """SELECT q.id AS question_id, q.text AS qtext, COALESCE(q.topic,'') AS topic,
                  t.title AS test_title,
                  COALESCE(s.title,'') AS subject,
                  COUNT(*) AS answered,
                  SUM(CASE WHEN aa.is_correct=1 THEN 0 ELSE 1 END) AS wrong,
                  COUNT(DISTINCT ta.user_id) AS students
           FROM attempt_answers aa
           JOIN test_attempts ta ON ta.id = aa.attempt_id
           JOIN questions q ON q.id = aa.question_id
           JOIN tests t ON t.id = ta.test_id
           LEFT JOIN lessons l ON l.test_id = t.id
           LEFT JOIN sections sec ON sec.id = l.section_id
           LEFT JOIN subjects s ON s.id = sec.subject_id
           WHERE COALESCE(aa.skipped,0)=0
           GROUP BY q.id
           HAVING wrong > 0
           ORDER BY wrong DESC, answered DESC""")
    items = []
    for r in rows[:limit]:
        r = dict(r)
        r["percent"] = round(r["wrong"] * 100 / max(1, r["answered"]))
        r["correct_ans"] = (db.fetchone(
            "SELECT text FROM question_options WHERE question_id=? AND is_correct=1 LIMIT 1",
            (r["question_id"],)) or {"text": ""})["text"]
        items.append(r)

    topics = {}
    for r in items:
        key = r["topic"] or "(без темы)"
        t = topics.setdefault(key, {"topic": key, "wrong": 0, "answered": 0})
        t["wrong"] += r["wrong"]
        t["answered"] += r["answered"]
    top_topics = sorted(topics.values(), key=lambda x: -x["wrong"])
    for t in top_topics:
        t["percent"] = round(t["wrong"] * 100 / max(1, t["answered"]))
    return {"items": items, "top_topics": top_topics[:15],
            "total_questions": len(items)}


async def admin_top_mistakes(request):
    from webapp import learning
    await learning._require_admin(request)
    data = await asyncio.to_thread(top_mistakes_sync)
    data.update(await auth.nav_context(request))
    return aiohttp_jinja2.render_template("admin_top_mistakes.html", request, data)


async def admin_top_mistakes_csv(request):
    from webapp import learning
    await learning._require_admin(request)
    data = await asyncio.to_thread(top_mistakes_sync, 5000)
    buf = io.StringIO()
    buf.write("\ufeff")
    w = csv.writer(buf, delimiter=";")
    w.writerow(["Место", "Ошиблись раз", "Ответов всего", "Учеников",
                "% ошибок", "Тема", "Предмет", "Тест", "Вопрос", "Правильный ответ"])
    for i, r in enumerate(data["items"], 1):
        w.writerow([i, r["wrong"], r["answered"], r["students"], r["percent"],
                    r["topic"], r["subject"], r["test_title"],
                    (r["qtext"] or "").replace("\n", " "), r["correct_ans"]])
    return web.Response(
        body=buf.getvalue().encode("utf-8"),
        headers={"Content-Type": "text/csv; charset=utf-8",
                 "Content-Disposition": 'attachment; filename="top_mistakes.csv"'})


def register_routes(app):
    app.router.add_get("/cabinet/errors/cards", mistakes_cards)
    app.router.add_get("/cabinet/errors/retest", mistakes_retest)
    app.router.add_get("/admin/learn/top-mistakes", admin_top_mistakes)
    app.router.add_get("/admin/learn/top-mistakes.csv", admin_top_mistakes_csv)
    app.router.add_get("/cabinet/errors", mistakes_page)
    app.router.add_get("/cabinet/errors/download.csv", mistakes_csv)
    app.router.add_get("/cabinet/errors/download.docx", mistakes_docx)
    app.router.add_get("/admin/learn/tests/{test_id:\\d+}/results", admin_test_results)
    app.router.add_get("/admin/learn/tests/{test_id:\\d+}/results.csv", admin_test_results_csv)
