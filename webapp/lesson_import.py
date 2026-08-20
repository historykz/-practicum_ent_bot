"""
Импорт тестов для уроков раздела "Начать обучение".

Переиспользует тот же текстовый формат, что уже привычен админам бота
(services/text_import_service.py, utils.parse_questions_text):

    Текст вопроса
    A) вариант
    B) вариант *   <- правильный
    C) вариант
    D) вариант

А также формат ZIP-архива из services/zip_import_service.py:
questions.txt с тегом [img:имя_файла.png] перед вопросом + папка images/.
Картинки для сайта сохраняются на диск (не в Telegram), рядом с базой
данных (Path(config.DB_PATH).parent/uploads/questions), чтобы пережить
деплой так же, как и сама база.

Импорт двухшаговый: сначала parse_draft_from_* (только парсит и
показывает предпросмотр, ничего не пишет в tests/questions), затем
после подтверждения администратором — finalize_test() создаёт
реальный тест в базе.
"""
import io
import logging
import re
import uuid
import zipfile
from pathlib import Path

import config
import database as db
from utils import parse_questions_text, _OPTION_RE, _LETTERS_LATIN, _LETTERS_CYR

logger = logging.getLogger(__name__)

_IMG_TAG_RE = re.compile(r"^\s*\[img:([^\]]+)\]\s*$", re.IGNORECASE)


def upload_dir() -> Path:
    d = Path(config.DB_PATH).resolve().parent / "uploads" / "questions"
    d.mkdir(parents=True, exist_ok=True)
    return d


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Текст конспекта из .docx: абзацы и ячейки таблиц по порядку."""
    import docx
    document = docx.Document(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts).strip()


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Текст конспекта из .pdf (только текстовый слой, без OCR сканов)."""
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(file_bytes))
    parts = [(page.extract_text() or "").strip() for page in reader.pages]
    return "\n\n".join(p for p in parts if p).strip()


def extract_lesson_content_from_upload(filename: str, file_bytes: bytes) -> str:
    """По расширению файла выбирает извлечение текста для конспекта урока."""
    ext = (filename or "").rsplit(".", 1)[-1].lower()
    if ext == "docx":
        return extract_text_from_docx(file_bytes)
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    raise ValueError(f"Неподдерживаемый формат файла: .{ext} (нужен .docx или .pdf)")


def _parse_with_images(raw: str) -> tuple[list[dict], list[str]]:
    """Как utils.parse_questions_text, но каждый блок может начинаться
    со строки [img:имя_файла] - тег сохраняется в question['image_filename']."""
    questions: list[dict] = []
    errors: list[str] = []

    lines = raw.replace("\r\n", "\n").split("\n")
    blocks: list[list[str]] = []
    current: list[str] = []
    for ln in lines:
        if ln.strip() == "":
            if current:
                blocks.append(current)
                current = []
        else:
            current.append(ln)
    if current:
        blocks.append(current)

    for bi, block in enumerate(blocks, start=1):
        image_filename = None
        if block and _IMG_TAG_RE.match(block[0]):
            image_filename = _IMG_TAG_RE.match(block[0]).group(1).strip()
            block = block[1:]

        if len(block) < 3:
            errors.append(f"Блок {bi}: слишком мало строк (нужен текст + минимум 2 варианта).")
            continue

        question_lines: list[str] = []
        option_lines: list[tuple[str, str, bool]] = []

        idx = 0
        while idx < len(block):
            line = block[idx]
            m = _OPTION_RE.match(line)
            if m:
                letter = m.group(1).upper()
                if letter in _LETTERS_LATIN or letter in _LETTERS_CYR:
                    break
            question_lines.append(line)
            idx += 1

        if not question_lines:
            errors.append(f"Блок {bi}: не найден текст вопроса.")
            continue

        correct_count = 0
        while idx < len(block):
            line = block[idx]
            idx += 1
            m = _OPTION_RE.match(line)
            if not m:
                if option_lines:
                    letter, txt, is_c = option_lines[-1]
                    option_lines[-1] = (letter, txt + " " + line.strip(), is_c)
                else:
                    question_lines.append(line)
                continue
            letter = m.group(1).upper()
            text = m.group(2).strip()
            is_correct = False
            if text.endswith("*"):
                is_correct = True
                text = text[:-1].rstrip()
                correct_count += 1
            option_lines.append((letter, text, is_correct))

        if len(option_lines) < 2:
            errors.append(f"Блок {bi}: меньше 2 вариантов ответа.")
            continue
        if len(option_lines) > 10:
            errors.append(f"Блок {bi}: больше 10 вариантов ответа.")
            continue
        if correct_count == 0:
            errors.append(f"Блок {bi}: не указан правильный ответ (символ *).")
            continue
        if correct_count > 1:
            errors.append(f"Блок {bi}: указано несколько правильных ответов.")
            continue

        question_text = "\n".join(question_lines).strip()
        if not question_text:
            errors.append(f"Блок {bi}: пустой текст вопроса.")
            continue

        correct_index = next(i for i, (_, _, c) in enumerate(option_lines) if c)
        questions.append({
            "text": question_text,
            "options": [opt[1] for opt in option_lines],
            "correct_index": correct_index,
            "image_filename": image_filename,
            "web_image_path": None,
        })

    return questions, errors


def parse_draft_from_text(raw_text: str) -> tuple[list[dict], list[str]]:
    """Только парсит текст (без картинок) — ничего не пишет в базу."""
    questions, errors = parse_questions_text(raw_text)
    for q in questions:
        q["web_image_path"] = None
    return questions, errors


def parse_draft_from_zip(zip_bytes: bytes) -> tuple[list[dict], list[str]]:
    """Парсит ZIP и СРАЗУ сохраняет картинки на диск (сам тест в базу
    ещё не пишет — только после подтверждения через finalize_test)."""
    try:
        zf = zipfile.ZipFile(io.BytesIO(zip_bytes))
    except zipfile.BadZipFile:
        return [], ["Файл не является ZIP-архивом."]

    txt_name = next((n for n in zf.namelist() if n.lower().endswith("questions.txt")), None)
    if not txt_name:
        txt_name = next((n for n in zf.namelist() if n.lower().endswith(".txt")), None)
    if not txt_name:
        return [], ["В архиве не найден questions.txt."]

    raw_text = zf.read(txt_name).decode("utf-8", errors="replace")
    images = {}
    for n in zf.namelist():
        base = n.rsplit("/", 1)[-1]
        if n.lower().startswith("images/") and base:
            images[base] = zf.read(n)

    questions, errors = _parse_with_images(raw_text)

    for q in questions:
        fname = q.get("image_filename")
        if fname and fname in images:
            ext = Path(fname).suffix or ".png"
            new_name = f"{uuid.uuid4().hex}{ext}"
            (upload_dir() / new_name).write_bytes(images[fname])
            q["web_image_path"] = f"/uploads/questions/{new_name}"

    return questions, errors


def finalize_test(title: str, created_by: int, questions: list[dict],
                   settings: dict = None) -> int:
    """Создаёт реальный тест + вопросы в базе из уже распарсенных данных
    (вызывается только после подтверждения администратором превью)."""
    settings = settings or {}
    db.execute(
        "INSERT INTO tests (title, created_by, status, test_type, language, "
        "show_correct, show_explanation, show_results, attempts_limit, "
        "time_per_question, shuffle_questions) "
        "VALUES (?, ?, 'active', 'regular', 'ru', ?, ?, ?, ?, ?, ?)",
        (
            title, created_by,
            1 if settings.get("show_correct", True) else 0,
            1 if settings.get("show_correct", True) else 0,
            1 if settings.get("show_results", True) else 0,
            int(settings.get("attempts_limit", 0) or 0),
            int(settings.get("time_per_question", 0) or 0),
            1 if settings.get("shuffle_questions", False) else 0,
        ),
    )
    test_id = db.fetchone("SELECT last_insert_rowid() AS id")["id"]

    for order, q in enumerate(questions):
        db.execute(
            "INSERT INTO questions (test_id, text, source_type, order_num, web_image_path) "
            "VALUES (?, ?, 'text_import', ?, ?)",
            (test_id, q["text"], order, q.get("web_image_path")),
        )
        qid = db.fetchone("SELECT last_insert_rowid() AS id")["id"]
        for oi, opt_text in enumerate(q["options"]):
            db.execute(
                "INSERT INTO question_options (question_id, text, is_correct, order_num) "
                "VALUES (?, ?, ?, ?)",
                (qid, opt_text, 1 if oi == q["correct_index"] else 0, oi),
            )
    return test_id


def delete_test(test_id: int) -> None:
    """Удаляет тест и все его вопросы/варианты (ON DELETE CASCADE)."""
    db.execute("DELETE FROM tests WHERE id=?", (test_id,))
